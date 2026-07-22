from __future__ import annotations

import io
import uuid
from pathlib import Path

from app.auth.dependencies import get_current_user
from app.database import get_session
from app.knowledge_base.config import (
    MAX_UPLOAD_MB,
    SUPPORTED_FORMATS,
    validate_upload,
)
from app.knowledge_base.manager import IngestionRejectedError, KnowledgeBaseManager
from app.knowledge_base.models import DocumentStatus, KnowledgeDocument
from app.knowledge_base.schemas import (
    DocumentResponse,
    DocumentUpdate,
    KnowledgeBaseConfigResponse,
)
from app.models import User
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import Response
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

router = APIRouter(prefix="/knowledge-base", tags=["Knowledge Base"])

# AUTH: business_id comes from the authenticated session (get_current_user),
# never from the client.

# Raw, original-format uploads are kept here so a document can later be
# previewed/downloaded in its original format, and so "Retry" can re-attempt
# learning without asking the business owner to upload the file again.
RAW_UPLOAD_DIR = Path("uploads/knowledge_base")

# Media types used for previewing ("View") a document in its original format.
_INLINE_MEDIA_TYPES = {
    ".pdf": "application/pdf",
    ".txt": "text/plain; charset=utf-8",
    ".md": "text/plain; charset=utf-8",
}
_DOWNLOAD_MEDIA_TYPES = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def extract_document_text(file_name: str, content: bytes) -> str:
    extension = Path(file_name).suffix.lower()

    if extension in {".md", ".txt"}:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError("Text files must use UTF-8 encoding.") from exc

    if extension == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    if extension == ".docx":
        from docx import Document

        document = Document(io.BytesIO(content))
        return "\n\n".join(paragraph.text for paragraph in document.paragraphs)

    raise ValueError("Unsupported document type.")


def _raw_path(business_id: uuid.UUID, stored_name: str) -> Path:
    return RAW_UPLOAD_DIR / str(business_id) / stored_name


async def _get_document_or_404(
    session: AsyncSession, business_id: uuid.UUID, document_id: uuid.UUID
) -> KnowledgeDocument:
    document = await session.scalar(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == document_id,
            KnowledgeDocument.business_id == business_id,
        )
    )
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found.")
    return document


async def _learn_document(
    manager: KnowledgeBaseManager,
    document: KnowledgeDocument,
    business_id: uuid.UUID,
    raw_content: bytes,
) -> None:
    """
    Extracts text from the raw uploaded bytes and teaches AISHA from it.

    Always leaves `document.status` as READY or FAILED and never raises —
    so a failed document still keeps its place in the list with a friendly,
    non-technical status the business owner can act on (Retry/Replace).
    """
    try:
        extracted_text = extract_document_text(document.file_name, raw_content)
    except ValueError as exc:
        document.status = DocumentStatus.FAILED
        document.error_message = str(exc)
        return
    except Exception:
        document.status = DocumentStatus.FAILED
        document.error_message = (
            "We couldn't read this document. Please try uploading it again."
        )
        return

    if not extracted_text.strip():
        document.status = DocumentStatus.FAILED
        document.error_message = "This document doesn't contain any readable text."
        return

    document_path = manager.resolver.resolve_within_tenant(
        business_id, document.stored_name
    )
    document_path.write_text(extracted_text, encoding="utf-8")

    try:
        await manager.ingest_document(business_id, document.stored_name)
    except IngestionRejectedError as exc:
        document.status = DocumentStatus.FAILED
        document.error_message = str(exc)
        return
    except Exception:
        document.status = DocumentStatus.FAILED
        document.error_message = (
            "Something went wrong while AISHA was learning from this document."
        )
        return

    document.status = DocumentStatus.READY
    document.error_message = None


@router.get("/config", response_model=KnowledgeBaseConfigResponse)
async def get_knowledge_base_config():
    """Exposes upload limits/formats so the frontend never hardcodes them."""
    return KnowledgeBaseConfigResponse(
        max_upload_size_mb=MAX_UPLOAD_MB,
        supported_formats=list(SUPPORTED_FORMATS),
    )


@router.get("/documents", response_model=list[DocumentResponse])
async def list_documents(
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """Returns every document for one business, newest first."""
    business_id = current_user.id
    await KnowledgeBaseManager.set_tenant_context(session, business_id)

    rows = await session.scalars(
        select(KnowledgeDocument)
        .where(KnowledgeDocument.business_id == business_id)
        .order_by(KnowledgeDocument.created_at.desc())
    )
    return rows.all()


@router.post(
    "/documents", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED
)
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """Uploads a new document and lets AISHA start learning from it."""
    business_id = current_user.id
    file_name = file.filename or ""
    safe_file_name = Path(file_name).name
    extension = Path(safe_file_name).suffix.lower()

    content = await file.read()
    validate_upload(safe_file_name, file_name, extension, content)

    manager = KnowledgeBaseManager(session)
    await manager.set_tenant_context(session, business_id)

    stored_name = f"{uuid.uuid4().hex}{extension}"
    raw_path = _raw_path(business_id, stored_name)
    raw_path.parent.mkdir(parents=True, exist_ok=True)
    raw_path.write_bytes(content)

    document = KnowledgeDocument(
        business_id=business_id,
        stored_name=stored_name,
        file_name=safe_file_name,
        display_name=safe_file_name,
        file_type=extension.lstrip(".").upper(),
        file_size=len(content),
        status=DocumentStatus.LEARNING,
    )
    session.add(document)
    await session.flush()

    await _learn_document(manager, document, business_id, content)

    await session.commit()
    await session.refresh(document)
    return document


@router.get("/documents/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """Returns the details of a single document (used by the View panel)."""
    business_id = current_user.id
    await KnowledgeBaseManager.set_tenant_context(session, business_id)
    return await _get_document_or_404(session, business_id, document_id)


@router.get("/documents/{document_id}/file")
async def get_document_file(
    document_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """
    Streams the original document file back for preview/download.

    PDF and text files are served inline so the browser can preview them;
    DOCX files are served as a download since browsers can't render them.
    """
    business_id = current_user.id
    await KnowledgeBaseManager.set_tenant_context(session, business_id)
    document = await _get_document_or_404(session, business_id, document_id)

    raw_path = _raw_path(business_id, document.stored_name)
    if not raw_path.is_file():
        raise HTTPException(
            status_code=404, detail="The original file is no longer available."
        )

    extension = Path(document.stored_name).suffix.lower()
    content = raw_path.read_bytes()

    if extension in _INLINE_MEDIA_TYPES:
        media_type = _INLINE_MEDIA_TYPES[extension]
        disposition = f'inline; filename="{document.file_name}"'
    else:
        media_type = _DOWNLOAD_MEDIA_TYPES.get(extension, "application/octet-stream")
        disposition = f'attachment; filename="{document.file_name}"'

    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": disposition},
    )


@router.patch("/documents/{document_id}", response_model=DocumentResponse)
async def update_document(
    document_id: uuid.UUID,
    updates: DocumentUpdate,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """
    Updates document metadata (display name, category, description, tags).

    This never touches the file or re-triggers learning — AISHA only
    relearns when the document content itself changes (see Replace).
    """
    business_id = current_user.id
    await KnowledgeBaseManager.set_tenant_context(session, business_id)
    document = await _get_document_or_404(session, business_id, document_id)

    if updates.display_name is not None:
        document.display_name = updates.display_name
    if updates.category is not None:
        document.category = updates.category or None
    if updates.description is not None:
        document.description = updates.description or None
    if updates.tags is not None:
        document.tags = updates.tags or None

    await session.commit()
    await session.refresh(document)
    return document


@router.put("/documents/{document_id}/replace", response_model=DocumentResponse)
async def replace_document(
    document_id: uuid.UUID,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """Replaces a document's file with a newer version, keeping the same record."""
    business_id = current_user.id
    file_name = file.filename or ""
    safe_file_name = Path(file_name).name
    extension = Path(safe_file_name).suffix.lower()

    content = await file.read()
    validate_upload(safe_file_name, file_name, extension, content)

    manager = KnowledgeBaseManager(session)
    await manager.set_tenant_context(session, business_id)
    document = await _get_document_or_404(session, business_id, document_id)

    # Keep the same stored_name so the new content updates (rather than
    # duplicates) the existing indexed knowledge for this document.
    raw_path = _raw_path(business_id, document.stored_name)
    raw_path.parent.mkdir(parents=True, exist_ok=True)
    raw_path.write_bytes(content)

    document.file_name = safe_file_name
    document.file_type = extension.lstrip(".").upper()
    document.file_size = len(content)
    document.status = DocumentStatus.LEARNING
    document.error_message = None
    await session.flush()

    await _learn_document(manager, document, business_id, content)

    await session.commit()
    await session.refresh(document)
    return document


@router.post("/documents/{document_id}/retry", response_model=DocumentResponse)
async def retry_document(
    document_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """Re-attempts learning for a document whose processing previously failed."""
    business_id = current_user.id

    manager = KnowledgeBaseManager(session)
    await manager.set_tenant_context(session, business_id)
    document = await _get_document_or_404(session, business_id, document_id)

    if document.status != DocumentStatus.FAILED:
        raise HTTPException(
            status_code=400, detail="Only failed documents can be retried."
        )

    raw_path = _raw_path(business_id, document.stored_name)
    if not raw_path.is_file():
        raise HTTPException(
            status_code=404,
            detail="The original file is no longer available. Please replace the document instead.",
        )

    document.status = DocumentStatus.LEARNING
    document.error_message = None
    await session.flush()

    await _learn_document(manager, document, business_id, raw_path.read_bytes())

    await session.commit()
    await session.refresh(document)
    return document


@router.delete("/documents/{document_id}", status_code=status.HTTP_200_OK)
async def delete_document(
    document_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """Deletes a document, its learned knowledge, and its stored file."""
    business_id = current_user.id

    manager = KnowledgeBaseManager(session)
    await manager.set_tenant_context(session, business_id)
    document = await _get_document_or_404(session, business_id, document_id)

    from app.knowledge_base.models import WikiChunk

    chunks = await session.scalars(
        select(WikiChunk).where(
            WikiChunk.business_id == business_id,
            WikiChunk.source_file == document.stored_name,
        )
    )
    for chunk in chunks.all():
        await session.delete(chunk)

    raw_path = _raw_path(business_id, document.stored_name)
    raw_path.unlink(missing_ok=True)

    try:
        clean_path = manager.resolver.resolve_within_tenant(
            business_id, document.stored_name
        )
        clean_path.unlink(missing_ok=True)
    except Exception:
        pass

    await session.delete(document)
    await session.commit()

    return {"message": "Document deleted successfully."}
